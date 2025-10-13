"""CV generation functions (PDF/LaTeX and Word)."""

import subprocess
import textwrap
import io

from src.config import MODEL, LATEX_TEMPLATE


def json_to_cv_pdf(client, cv_dict):
    """Generates a CV PDF from the JSON data.

    Args:
        client: OpenAI client instance
        cv_dict: Dictionary containing CV data

    Returns:
        bool: True if PDF was generated successfully
    """
    prompt = textwrap.dedent(f"""
        - Lag en proffesjonell CV i latex format basert på JSON data.
        - Bruk informasjon som alder og erfaringer til å tilpasse CVen.
        - Rydd opp i dataen, du kan omformulere og fjerne punkter som er irrelevante, forvirrende eller duplikater.
        - Resultatet skal være en ferdig CV leveringsklar til arbeidsgiver (fjern tomme seksjoner).
        - JSON data: {cv_dict}
        - Bruk denne latex malen: {LATEX_TEMPLATE}
        """)

    latex_response_gen = client.chat.completions.create(
        model=MODEL,
        messages=[
            {"role": "system", "content": "You are a LaTeX generator."},
            {"role": "user", "content": prompt}
        ],
    )
    latex_cv_code = latex_response_gen.choices[0].message.content

    filename = "CV.tex"
    with open(filename, mode="w", encoding="utf-8") as f:
        f.write(latex_cv_code)

    # Compile to PDF using pdflatex
    subprocess.run(["pdflatex", "-interaction=nonstopmode", filename])

    return True


def generate_word_docx(client, cv_dict):
    """Generates a CV Word document from the JSON data.

    Args:
        client: OpenAI client instance
        cv_dict: Dictionary containing CV data

    Returns:
        BytesIO: Buffer containing the Word document
    """
    # Lazy import to avoid requiring docx if not used
    from docx import Document
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    from docx.shared import Pt

    # Get LLM to clean up the CV data
    prompt = textwrap.dedent(f"""
        - Rydd opp i denne CV dataen, du kan omformulere og fjerne punkter som er irrelevante, forvirrende eller duplikater.
        - Returner KUN den oppdaterte CV dataen i samme JSON format som du fikk.
        - Sørg for at all dataen er optimalt formatert for en CV.
        - Din output data skal brukes direkte til å generere en CV i Word docx.
        - CV data: {cv_dict}
        """)

    json_response_gen = client.chat.completions.create(
        model=MODEL,
        messages=[
            {"role": "system", "content": "You are a CV data cleaner."},
            {"role": "user", "content": prompt}
        ],
    )
    json_response_str = json_response_gen.choices[0].message.content

    try:
        import json
        cv_dict = json.loads(json_response_str)
    except json.JSONDecodeError as e:
        print(f"Error decoding JSON: {e}")
        return None

    # Create Word document
    doc = Document()
    personalia = cv_dict.get("Personalia", {})

    # Title with name
    title = doc.add_paragraph()
    title_run = title.add_run(personalia.get("Navn", ""))
    title_run.bold = True
    title_run.font.size = Pt(24)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Personalia details
    for key, value in personalia.items():
        if key == "Navn" or not value:
            continue
        p = doc.add_paragraph()
        run = p.add_run(f"{key}: ")
        run.bold = True
        p.add_run(str(value))

    doc.add_paragraph()

    # Education
    utdanning = cv_dict.get("Utdanning", [])
    if utdanning:
        doc.add_heading("Utdanning", level=1)
        for entry in utdanning:
            summary = f"{entry.get('Grad', '')} – {entry.get('Skole', '')} ({entry.get('Trinn/ferdig_år', '')})"
            doc.add_paragraph(summary, style="List Bullet")
            if entry.get("Ytterligere_informasjon"):
                doc.add_paragraph(entry["Ytterligere_informasjon"], style="List Continue")

    # Work experience
    arbeid = cv_dict.get("Arbeidserfaring", {})
    if arbeid.get("Stillinger") or arbeid.get("Dugnad"):
        doc.add_heading("Arbeidserfaring", level=1)
        for stilling in arbeid.get("Stillinger", []):
            bullet = doc.add_paragraph(style="List Bullet")
            header = bullet.add_run(
                f"{stilling.get('Tittel', '')}, {stilling.get('Firma', '')}, {stilling.get('Periode', '')}"
            )
            header.bold = True
            if stilling.get("Beskrivelse"):
                doc.add_paragraph(stilling["Beskrivelse"], style="List Continue")

        if arbeid.get("Dugnad"):
            doc.add_paragraph().add_run("Dugnad").italic = True
            for dugnad in arbeid["Dugnad"]:
                doc.add_paragraph(
                    f"{dugnad.get('Oppdrag', '')} ({dugnad.get('Periode', '')})",
                    style="List Bullet"
                )
                if dugnad.get("Beskrivelse"):
                    doc.add_paragraph(dugnad["Beskrivelse"], style="List Continue")

    # Skills
    ferdigheter = cv_dict.get("Ferdigheter", {})
    if any(ferdigheter.values()):
        doc.add_heading("Ferdigheter", level=1)
        for skill in ferdigheter.get("Ferdigheter_og_kompetanser", []):
            doc.add_paragraph(
                f"{skill.get('Ferdighet', '')} – {skill.get('Nivå', '')}",
                style="List Bullet"
            )
            if skill.get("Beskrivelse"):
                doc.add_paragraph(skill["Beskrivelse"], style="List Continue")

        språk = ferdigheter.get("Språk", [])
        if språk:
            doc.add_paragraph("Språk:", style="Heading 3")
            for item in språk:
                doc.add_paragraph(
                    f"{item.get('Språk', '')}: {item.get('Nivå', '')}",
                    style="List Bullet"
                )

        sertifikater = ferdigheter.get("Sertifikater", [])
        if sertifikater:
            doc.add_paragraph("Sertifikater:", style="Heading 3")
            for cert in sertifikater:
                doc.add_paragraph(cert, style="List Bullet")

        annet = ferdigheter.get("Annet", [])
        if annet:
            doc.add_paragraph("Annet:", style="Heading 3")
            for entry in annet:
                doc.add_paragraph(entry, style="List Bullet")

    # Interests
    interesser = cv_dict.get("Interesser_og_hobbyer", [])
    if interesser:
        doc.add_heading("Interesser og hobbyer", level=1)
        for entry in interesser:
            doc.add_paragraph(entry.get("Interesse/Hobby", ""), style="List Bullet")
            if entry.get("Beskrivelse"):
                doc.add_paragraph(entry["Beskrivelse"], style="List Continue")

    # Future goals
    mål = cv_dict.get("Fremtidige_mål", {})
    if mål:
        doc.add_heading("Fremtidige mål", level=1)
        if mål.get("Fremtidsutsikter_og_mål"):
            doc.add_paragraph(mål["Fremtidsutsikter_og_mål"])
        if mål.get("Jobbønsker"):
            doc.add_paragraph("Jobbønsker:", style="Heading 3")
            for entry in mål["Jobbønsker"]:
                doc.add_paragraph(entry.get("Jobbønske", ""), style="List Bullet")
                if entry.get("Begrunnelse"):
                    doc.add_paragraph(entry["Begrunnelse"], style="List Continue")

    # Save to buffer
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer
